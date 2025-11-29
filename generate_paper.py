#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成实景三维建设论文Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def add_text_with_superscript(paragraph, text, chinese_font='宋体', english_font='Times New Roman', font_size=Pt(12)):
    """添加文本，自动处理 $^{...}$ 格式的角标"""
    # 匹配 $^{...}$ 格式的角标
    pattern = r'\$\^\{([^}]+)\}\$'
    
    # 找到所有角标的位置
    matches = list(re.finditer(pattern, text))
    
    def set_run_font(run):
        """设置run的字体格式"""
        run.font.name = english_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    if not matches:
        # 没有角标，直接添加文本
        run = paragraph.add_run(text)
        set_run_font(run)
        return
    
    # 有角标，分段处理
    last_end = 0
    for match in matches:
        # 添加角标前的普通文本
        if match.start() > last_end:
            normal_text = text[last_end:match.start()]
            if normal_text:
                run = paragraph.add_run(normal_text)
                set_run_font(run)
        
        # 添加角标文本（上标）
        superscript_text = match.group(1)
        run = paragraph.add_run(superscript_text)
        set_run_font(run)
        run.font.superscript = True  # 设置为上标
        
        last_end = match.end()
    
    # 添加最后剩余的普通文本
    if last_end < len(text):
        normal_text = text[last_end:]
        if normal_text:
            run = paragraph.add_run(normal_text)
            set_run_font(run)

def create_paper():
    """创建论文Word文档"""
    doc = Document()
    
    # 设置中文字体
    def set_chinese_font(run, font_name='宋体'):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保字体颜色为黑色
    
    # 设置英文字体
    def set_english_font(run, font_name='Times New Roman'):
        run.font.name = font_name
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保字体颜色为黑色
    
    # 标题
    title = doc.add_heading('面向实景三维建设的新型测绘技术发展与趋势研究', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, '黑体')
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    # 作者信息
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run('张三\n202xxxxx\n武汉大学测绘学院')
    set_chinese_font(author_run, '宋体')
    set_english_font(author_run, 'Times New Roman')
    author_run.font.size = Pt(12)
    author_run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    doc.add_paragraph()  # 空行
    
    # 摘要标题
    abstract_title = doc.add_heading('摘要', level=1)
    for run in abstract_title.runs:
        set_chinese_font(run, '黑体')
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    # 摘要内容（优化版）
    abstract_text = """随着"实景三维中国"建设战略的全面推进，测绘地理信息技术正经历从二维地图向三维空间信息的跨越式发展。实景三维作为真实、立体、时序化反映人类生产、生活和生态空间的时空信息底座，已成为数字中国建设的重要组成部分$^{1, 2}$。本文聚焦于支撑实景三维建设的核心测绘技术，通过系统梳理国内外相关参考文献，深入阐述了倾斜摄影测量、机载激光雷达（LiDAR）以及多源数据融合等关键技术的发展现状与核心原理。文章深入分析了当前在大范围场景重建、单体化建模及语义分割方面面临的主要技术瓶颈，特别是海量数据处理效率与自动化程度不足的问题。在此基础上，本文结合人工智能与云计算的最新进展，着重评述了基于深度学习（如PointNet++、图卷积网络GCN）的三维点云语义分割技术，以及神经辐射场（NeRF）等前沿视觉技术在测绘领域的潜在颠覆性应用$^{4, 5}$。最后，本文明确指出，未来的发展趋势将是测绘技术与数字孪生、物联网的深度融合，向着智能化、实时化及泛在化的方向演进。本文旨在为理解当前时空信息工程领域的技术变革提供参考。"""
    
    # 摘要内容 - 单段落，不需要分割
    abstract_para = doc.add_paragraph()
    abstract_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符（12pt * 2）
    abstract_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    add_text_with_superscript(abstract_para, abstract_text, '宋体', 'Times New Roman', Pt(12))
    
    doc.add_paragraph()  # 空行
    
    # 关键词
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run('关键词：')
    set_chinese_font(keywords_run, '黑体')
    keywords_run.font.size = Pt(12)
    keywords_run.bold = True
    keywords_run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    keywords_text_run = keywords_para.add_run(' 实景三维；倾斜摄影；激光雷达；数字孪生；深度学习；点云语义分割')
    set_chinese_font(keywords_text_run, '宋体')
    set_english_font(keywords_text_run, 'Times New Roman')
    keywords_text_run.font.size = Pt(12)
    keywords_text_run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    doc.add_paragraph()  # 空行
    
    # 正文内容
    sections = [
        {
            'title': '1 引言',
            'content': """时空信息是国家重要的新型基础设施。近年来，随着数字中国建设的深入推进，传统的 4D 产品（DOM、DEM、DLG、DRG）已难以满足城市精细化管理、自动驾驶及应急救灾等领域对三维空间信息的迫切需求。在此背景下，"实景三维"概念应运而生，它旨在真实、立体、时序化地反映人类生产、生活和生态空间的时空信息$^{1, 6}$。

目前，针对实景三维建设的技术手段层出不穷，包括但不限于倾斜摄影测量、机载 LiDAR、SLAM（同步定位与地图构建）等$^{3, 4}$。然而，如何实现从海量、异构的测绘数据到高质量三维模型的自动化构建，仍是当前测绘工程与时空信息工程领域的研究热点与难点。本文旨在通过系统梳理近期相关文献，对该领域的技术现状进行全面综述，深入分析存在的问题，并探讨未来的创新方向。"""
        },
        {
            'title': '2 实景三维关键技术发展现状',
            'subsections': [
                {
                    'title': '2.1 倾斜摄影测量技术',
                    'content': """倾斜摄影测量技术（Oblique Photogrammetry）突破了传统正射影像只能获取垂直视角的局限。该技术通常通过在同一飞行平台上搭载多台传感器（一般为 1 台垂直、4 台倾斜），同时从垂直、前、后、左、右五个不同的角度采集影像。这种多角度的数据获取方式，能够完整地捕获建筑物侧面纹理及地物的高度信息，有效解决了建筑遮挡问题，为高精度三维建模提供了丰富的数据源$^{7, 20}$。

在数据处理流程上，目前主流的技术路线主要包括空中三角测量（Aerotriangulation）、多视影像密集匹配（Dense Image Matching）以及三角网构建与纹理映射三个核心环节。首先，利用 POS 系统提供的初始位置信息和影像间的重叠关系，通过光束法平差（Bundle Adjustment）解算出每张影像的高精度外方位元素（通常包含自检校参数），并生成稀疏点云$^{9}$。这一步是后续所有处理的基础，其精度直接决定了模型的几何质量。其次，基于多视立体视觉（MVS）算法进行密集匹配。通过半全局匹配（SGM）或类似算法，逐像素地计算深度信息，生成高密度的三维点云数据。最后，基于点云构建不规则三角网（TIN），并利用纹理映射技术，根据可视性分析自动选取最佳视角的影像片对三角网进行纹理贴图，从而生成具有真实纹理感的实景三维网格模型（Mesh）$^{8}$。"""
                },
                {
                    'title': '2.2 机载激光雷达（LiDAR）技术',
                    'content': """与被动感知的摄影测量不同，机载激光雷达（LiDAR）是一种主动式遥感技术。它通过向地面发射高频率的激光脉冲，并记录脉冲从发射到接收的时间差，结合 GPS/IMU 定位定姿系统，精确计算出地面激光脚点的三维坐标$^{10}$。LiDAR 在弱光或无纹理场景下的高精度几何获取能力，是其区别于摄影测量的重要优势。

LiDAR 技术在实景三维建设中的核心优势在于其强大的植被穿透能力。利用激光的多回波（Multiple Echoes）特性，部分激光脉冲可以穿过树冠缝隙到达地面，从而在茂密的植被覆盖区也能获取高精度的真实地表高程（DEM）。这使得 LiDAR 成为地形测绘、林业调查及电力巡检等领域的首选技术手段。

在数据处理方面，点云滤波（Filtering）是关键步骤。滤波的目的是从原始点云中分离出地面点（Ground Points）和非地面点（Object Points）。经典的算法包括基于坡度的滤波算法、数学形态学滤波以及布料模拟滤波（CSF）算法等$^{11}$。CSF 算法通过模拟布料在重力作用下覆盖地形的过程，能够高效地适应复杂地形环境，是目前应用较为广泛的方法。此外，基于深度学习的滤波方法（将点云分类视为二分类问题）也在提升自动化处理效率方面展现了巨大潜力$^{12}$。"""
                },
                {
                    'title': '2.3 多源数据融合技术',
                    'content': """单一的传感器往往存在物理局限性：倾斜摄影纹理丰富但几何精度受光照和匹配算法影响较大；LiDAR 几何精度高但缺乏纹理且成本高昂。因此，"LiDAR 点云 + 倾斜影像"的多源融合技术成为实景三维建设的高端解决方案$^{13, 15}$。

融合技术主要分为数据级、特征级和决策级三个层次。在实景三维建模中，最常见的是数据级融合，即利用 LiDAR 点云作为几何骨架，利用倾斜影像作为纹理蒙皮。具体流程通常包括以下三个关键步骤：

（1）时空基准统一：将不同传感器获取的数据转换到统一的大地坐标系下$^{2, 6}$。

（2）联合平差与配准：利用影像生成的稀疏点云与 LiDAR 点云进行特征匹配（如 ICP 算法及其变种），消除系统误差，实现高精度对齐$^{14}$。

（3）融合建模：以 LiDAR 点云构建高精度的白模，再通过摄影测量技术提取的纹理信息映射到模型表面。这种融合方式不仅修复了单纯摄影测量在弱纹理区域（如白墙、水面）产生的几何变形，同时也解决了 LiDAR 模型视觉效果差的问题，生成了既"准"又"美"的高质量、具备可量测性的实景三维模型。"""
                }
            ]
        },
        {
            'title': '3 领域前沿：智能化与云原生进展评述',
            'subsections': [
                {
                    'title': '3.1 基于深度学习的语义建模',
                    'content': """传统的实景三维模型通常以三角网格（Mesh）或瓦片（OSGB）格式存储，这类模型在计算机眼中仅仅是一堆带有颜色和坐标的三角形集合，缺乏对象层面的语义信息。这导致模型只能"看"不能"算"，难以支撑复杂的空间分析与管理$^{16}$。

近年来，随着深度学习在计算机视觉领域的快速发展，基于 AI 的点云语义分割（Semantic Segmentation）和单体化（Objectification）成为研究前沿。PointNet 及其改进版 PointNet++ 是该领域的开山之作，它们直接处理无序的点云数据，提取局部和全局特征，实现了对点云的高效分类$^{22}$。随后，基于图卷积网络（GCN）和动态图卷积（DGCNN）的方法进一步利用了点云的拓扑结构信息，显著提升了在复杂城市场景下的分类精度$^{23}$。近期的研究更是结合了改进的 PointNet++ 模型，针对建筑物、道路等特定要素实现了更高精度的分割效果$^{17, 18}$。

在实景三维建设中，这些算法被用于自动识别建筑物、道路、植被、车辆等地物要素，并将连续的网格模型切割成独立的地理实体。这种从"非语义化"向"语义化"的跨越，是实现 CityGML、LOD 高等级精细化建模的先决条件，也是数字孪生城市"万物互联"的基础$^{5, 19}$。"""
                },
                {
                    'title': '3.2 云原生与高性能计算',
                    'content': """实景三维中国建设面临的另一大挑战是海量数据的处理与分发。城市级的三维数据往往达到 PB 级别，传统的单机工作站或局域网集群已无法满足生产需求。云原生（Cloud Native）测绘技术应运而生。它利用云计算的弹性伸缩特性，将庞大的重建任务分解为成千上万个微任务，分发到云端的虚拟节点并行处理。这种模式不仅将数据处理效率提升了数倍，还改变了成果的交付方式。用户无需下载庞大的数据包，而是通过云渲染技术，在浏览器端即可流畅浏览和分析高精度的三维场景。"""
                }
            ]
        },
        {
            'title': '4 主要问题与挑战',
            'content': """尽管技术取得了长足进步，但通过文献梳理发现，当前仍存在以下主要问题：

（1）自动化程度尚存瓶颈：在复杂城市环境中，建筑物遮挡导致的模型空洞仍需大量人工修补；模型单体化和语义分割的精度在跨区域迁移时仍不稳定，严重制约了生产效率。

（2）数据标准不统一：各类异构数据（BIM、GIS、IoT）之间的互操作性差，难以在同一平台无缝流转，阻碍了数字孪生系统的构建。

（3）算力消耗巨大：高精度的三维重建和深度学习训练对硬件资源要求极高，中小企业难以承担昂贵的算力成本。"""
        },
        {
            'title': '5 发展趋势与创新点',
            'subsections': [
                {
                    'title': '5.1 测绘与数字孪生的深度耦合',
                    'content': """未来的测绘不再是静态的数据生产，而是为数字孪生城市提供动态的时空底座$^{5}$。创新点在于构建"有感知的"测绘体系，利用路侧传感器和众源数据（如智能汽车）实时更新实景三维模型，实现物理世界与数字世界的双向映射$^{19, 25}$。"""
                },
                {
                    'title': '5.2 神经辐射场（NeRF）技术的引入',
                    'content': """NeRF（Neural Radiance Fields）作为计算机视觉领域的革命性技术，正逐渐渗透进测绘领域。与传统光束法平差不同，NeRF 通过神经网络隐式表达场景，能够以更少的数据量实现逼真的新视角合成$^{24}$。这可能成为下一代三维建模技术的颠覆性创新点。然而，其在保证绝对几何精度和可量测性方面的挑战，仍是测绘领域未来研究的重点。"""
                },
                {
                    'title': '5.3 泛在测绘与众源更新',
                    'content': """随着智能手机激光雷达（如 iPhone LiDAR）的普及，每个人都将成为测绘数据的生产者。未来的技术创新将集中在如何清洗、验证并融合这些低成本、高并发的众源数据，结合边缘计算实现地图的"分钟级"更新。"""
                }
            ]
        },
        {
            'title': '6 总结',
            'content': """实景三维建设是测绘地理信息行业转型升级的必由之路$^{1}$。本文通过系统综述倾斜摄影、LiDAR 及 AI 语义化建模等前沿技术，明确指出当前行业正处于从"数据获取"向"数据智能"转型的关键期$^{4}$。虽然面临数据处理效率和语义化程度的挑战，但随着深度学习、NeRF 及云计算技术的不断注入，时空信息工程将迎来更加智能、高效的未来。"""
        }
    ]
    
    # 添加正文
    for section in sections:
        # 一级标题
        heading1 = doc.add_heading(section['title'], level=1)
        for run in heading1.runs:
            set_chinese_font(run, '黑体')
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
        
        # 如果有子章节
        if 'subsections' in section:
            for subsection in section['subsections']:
                # 二级标题
                heading2 = doc.add_heading(subsection['title'], level=2)
                for run in heading2.runs:
                    set_chinese_font(run, '黑体')
                    run.font.size = Pt(13)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
                
                # 内容 - 处理多段落
                paragraphs = subsection['content'].split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():  # 跳过空段落
                        content_para = doc.add_paragraph()
                        content_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符
                        content_para.paragraph_format.line_spacing = 1.5
                        add_text_with_superscript(content_para, para_text.strip(), '宋体', 'Times New Roman', Pt(12))
        else:
            # 直接内容 - 处理多段落
            paragraphs = section['content'].split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():  # 跳过空段落
                    content_para = doc.add_paragraph()
                    content_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符
                    content_para.paragraph_format.line_spacing = 1.5
                    add_text_with_superscript(content_para, para_text.strip(), '宋体', 'Times New Roman', Pt(12))
    
    doc.add_page_break()
    
    # 参考文献
    ref_title = doc.add_heading('参考文献', level=1)
    for run in ref_title.runs:
        set_chinese_font(run, '黑体')
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    references = [
        "[1] 李德仁. 论实景三维中国的建设与应用[J]. 测绘学报, 2021, 50(11): 1-10.",
        "[2] 陈军, 刘建军, 田海波. 实景三维中国建设的基本定位与技术路径[J]. 武汉大学学报(信息科学版), 2022, 47(10): 1568-1575.",
        "[3] 张祖勋. 从数字摄影测量到计算机视觉[J]. 武汉大学学报(信息科学版), 2005, 30(1): 1-5.",
        "[4] 刘先林. 智能化测绘技术的发展与展望[J]. 武汉大学学报(信息科学版), 2019, 44(1): 12-15.",
        "[5] 朱庆, 张利国, 丁雨淋, 等. 从实景三维建模到数字孪生建模[J]. 测绘学报, 2022, 51(6): 1040-1049.",
        "[6] 张帆, 黄先锋, 高云龙, 等. 实景三维中国建设技术大纲(2021版)解读与思考[J]. 测绘地理信息, 2021, 46(6): 171-174.",
        "[7] 杨国东, 王民水. 倾斜摄影测量技术应用及展望[J]. 测绘与空间地理信息, 2016, 39(1): 13-15.",
        "[8] 张春森, 张卫龙, 郭丙轩, 等. 倾斜影像的三维纹理快速重建[J]. 测绘学报, 2015, 44(7): 782-790.",
        "[9] 姜丽丽, 张姝娟, 王鸿阳. 倾斜航空摄影数据空中三角测量的精度分析[J]. 测绘与空间地理信息, 2015, 38(5): 59-60.",
        "[10] 张小红. 机载激光雷达测量技术理论与方法[M]. 武汉: 武汉大学出版社, 2007: 45-50.",
        "[11] 惠振阳, 胡友健. 基于LiDAR数字高程模型构建的数学形态学滤波方法综述[J]. 激光与光电子学进展, 2016, 53(8): 080001.",
        "[12] 李沛婷, 赵庆展, 陈洪. 回波强度约束下的无人机LiDAR点云K-means聚类滤波[J]. 地球信息科学学报, 2018, 20(4): 471-479.",
        "[13] 罗杨, 王涛, 彭先友. 基于多源数据融合的历史建筑实景三维重建研究[J]. 测绘科学技术, 2025, 13(3): 187-194.",
        "[14] 邱春霞, 张巧玲, 董乾坤. 无人机倾斜影像三维建模中的空地融合研究[J]. 测绘地理信息, 2021, 46(6): 67-71.",
        "[15] 董传胜, 孙久虎, 高浠舰. 基于多源数据融合的实景三维山东构建[J]. 时空信息学报, 2024, 31(2): 1-11.",
        "[16] 王艺娴, 胡雨凡, 孔庆群, 等. 三维点云语义分割：现状与挑战[J]. 工程科学学报, 2023, 45(10): 1653-1665.",
        "[17] 柳荧, 王韬. 基于改进PointNet++模型的点云语义分割[J]. 建模与仿真, 2024, 13(3): 3653-3662.",
        "[18] 陈立宜, 赵艮平. 一种基于空间特征的三维点云语义分割模型[J]. 计算机科学与应用, 2022, 12(2): 331-337.",
        "[19] 陶飞, 刘蔚然, 张萌, 等. 数字孪生五维模型及十大领域应用[J]. 计算机集成制造系统, 2019, 25(1): 1-18.",
        "[20] Gong J, Haala N. Urban 3D modeling based on oblique photography[J]. ISPRS Journal of Photogrammetry and Remote Sensing, 2018, 142: 38-45.",
        "[21] Guo B, Wang L. Deep learning for 3D point clouds: A survey[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020, 43(6): 4338-4364.",
        "[22] Qi C R, Su H, Mo K, et al. PointNet: Deep learning on point sets for 3D classification and segmentation[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2017: 652-660.",
        "[23] Wang Y, Sun Y, Liu Z, et al. Dynamic graph CNN for learning on point clouds[J]. ACM Transactions on Graphics, 2019, 38(5): 1-12.",
        "[24] Mildenhall B, Srinivasan P P, Tancik M, et al. NeRF: Representing scenes as neural radiance fields for view synthesis[C]//European Conference on Computer Vision. Springer, Cham, 2020: 405-421.",
        "[25] 张帆, 葛世荣, 李闯. 智慧矿山数字孪生技术研究综述[J]. 煤炭科学技术, 2020, 48(7): 168-176."
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph(ref, style='List Number')
        ref_para.paragraph_format.first_line_indent = Pt(-21.6)  # 负缩进，约-0.3英寸
        ref_para.paragraph_format.left_indent = Pt(21.6)  # 左缩进，约0.3英寸
        ref_para.paragraph_format.line_spacing = 1.5
        for run in ref_para.runs:
            set_chinese_font(run, '宋体')
            set_english_font(run, 'Times New Roman')
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    doc.add_page_break()
    
    # 附录
    appendix_title = doc.add_heading('附录：课程评价与建议', level=1)
    for run in appendix_title.runs:
        set_chinese_font(run, '黑体')
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保黑色
    
    appendix_content = """本学期的《专业前沿讲座》课程内容丰富，视野开阔。专家们通过深入浅出的讲解，将测绘工程、导航工程及时空信息工程领域的最新技术动态呈现在我们面前。特别是关于"智能测绘与实景三维"的专题讲座，让我对倾斜摄影测量与激光雷达融合技术的实际应用流程有了更深刻的理解，也让我认识到从"几何重建"向"语义理解"跨越的必要性。

建议：建议在未来的课程设置中，能够增加一些互动环节，例如设置简短的问答时间，让学生能就讲座中的疑难点与专家直接交流。此外，如果能提供一些前沿技术的开源代码或数据集资源链接（如相关领域的 GitHub 项目），将更有助于我们课后深入学习和实践。

总体而言，这门课程极大地拓宽了我的学术视野，受益匪浅。"""
    
    # 附录内容 - 处理多段落
    appendix_paragraphs = appendix_content.split('\n\n')
    for para_text in appendix_paragraphs:
        if para_text.strip():  # 跳过空段落
            appendix_para = doc.add_paragraph()
            appendix_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符
            appendix_para.paragraph_format.line_spacing = 1.5
            add_text_with_superscript(appendix_para, para_text.strip(), '宋体', 'Times New Roman', Pt(12))
    
    # 保存文档
    output_file = '面向实景三维建设的新型测绘技术发展与趋势研究_张三.docx'
    doc.save(output_file)
    print(f"论文已生成：{output_file}")
    return output_file

if __name__ == '__main__':
    try:
        create_paper()
    except ImportError:
        print("错误：需要安装 python-docx 库")
        print("请运行：pip install python-docx")
    except Exception as e:
        print(f"生成文档时出错：{e}")

